import docx
import re
import logging
from pathlib import Path
from docx.text.run import Run
from file_handlers.base_file_handler import BaseFileHandler


class DOCXFileHandler(BaseFileHandler):
    """Class for handling DOCX files."""

    TAG_PREFIX = "<RUN"
    TAG_SUFFIX = "/>"

    def _generate_tag(self, index: int) -> str:
        """Generate a unique tag for a text fragment."""
        return f"{self.TAG_PREFIX}{index}{self.TAG_SUFFIX}"

    def _get_run_formatting_signature(self, run: Run) -> tuple[any, ...]:
        """
        Creates a "fingerprint" of the run's formatting for comparison.
        If merging isn't enabled, uses the run's unique id(so no merging).
        """
        if not self.config["processing"]["docx_merge_runs"]:
            return tuple([id(run)])
        font = run.font
        return (
            run.bold,
            run.italic,
            run.underline,
            font.name,
            font.size,
            font.color.rgb if font.color and font.color.rgb else None
        )

    def extract_text(self, filepath: Path) -> str:
        """Extract text from DOCX, inserting tags before each text fragment, and returns a string."""
        try:
            document = docx.Document(str(filepath))
            text_parts = []
            run_index = 0

            for para in document.paragraphs:
                current_signature = None
                for run in para.runs:
                    if not current_signature or current_signature != self._get_run_formatting_signature(run):
                        # Starting new run
                        current_signature = self._get_run_formatting_signature(run)
                        run_index += 1
                        tag = self._generate_tag(run_index)
                        text_parts.append(tag)
                    text_parts.append(run.text)


            full_text = "".join(text_parts)
            return full_text

        except Exception as e:
            logging.error(f"Error extracting text/mapping from {filepath}: {e}")
            return ""

    def insert_text(self, original_filepath: Path, processed_text: str, output_filepath: Path):
        """Update the text of fragments in the document based on the processed string with tags."""
        try:
            document = docx.Document(str(original_filepath))
            new_run_map: dict[str, Run] = {}
            leftover_runs_map: dict[str, list[Run]] = {}
            run_index = 0

            for para in document.paragraphs:
                current_signature = None
                for run in para.runs:
                    if not current_signature or current_signature != self._get_run_formatting_signature(run):
                        # Starting new run
                        current_signature = self._get_run_formatting_signature(run)
                        run_index += 1
                        tag = self._generate_tag(run_index)
                        new_run_map[tag] = run
                    else:
                        tag = self._generate_tag(run_index)
                        if not tag in leftover_runs_map:
                            leftover_runs_map[tag] = []
                        leftover_runs_map[tag].append(run)

            tag_pattern = re.compile(f"({self.TAG_PREFIX}\\d+{self.TAG_SUFFIX})", re.DOTALL)

            parts = tag_pattern.split(processed_text)

            processed_runs_count = 0
            missed_tags = set(new_run_map.keys())

            for i in range(1, len(parts), 2):
                tag = parts[i]
                text_after_tag = parts[i+1] if (i+1) < len(parts) else ""

                if tag in new_run_map:
                    run_object = new_run_map[tag]
                    run_object.text = text_after_tag
                    processed_runs_count += 1
                    missed_tags.discard(tag)
                else:
                    logging.warning(f"Tag '{tag}' found in LLM response but not in original document map. Skipping.")

            if missed_tags:
                logging.warning(f"{len(missed_tags)} original tags were not found in the LLM response. Their corresponding run texts were not updated.")
                logging.warning(f"Missing tags: {missed_tags}")

            # Clearing leftover runs whose text has been transferred to another run.
            for tag in leftover_runs_map:
                for run in leftover_runs_map[tag]:
                    run.text = ""

            document.save(str(output_filepath))

        except Exception as e:
            logging.error(f"Error inserting processed text into {output_filepath}: {e}")